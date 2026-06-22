import os
import random
import string
from faker import Faker
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
import zipfile
import rarfile # Требует установки unrar
import py7zr

fake = Faker('ru_RU')

# Создаем корневую папку для тестовых данных
storage_path = 'test_storage'
if not os.path.exists(storage_path):
    os.makedirs(storage_path)

def generate_text_file(path, content=None):
    """Генерирует простой текстовый файл."""
    if content is None:
        content = fake.text()
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)

def generate_docx_file(path):
    """Генерирует .docx файл с абзацами."""
    doc = Document()
    doc.add_heading(fake.sentence(), level=1)
    for _ in range(random.randint(2, 5)):
        doc.add_paragraph(fake.paragraph())
    doc.save(path)

def generate_xlsx_file(path):
    """Генерирует .xlsx файл с таблицей."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Данные"
    # Заголовки
    ws.append(["Имя", "Фамилия", "Email", "Телефон"])
    for _ in range(10):
        ws.append([fake.first_name(), fake.last_name(), fake.email(), fake.phone_number()])
    wb.save(path)

def generate_pdf_file(path):
    """Генерирует простой .pdf файл."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(200, 10, txt=fake.text())
    pdf.output(path)

def create_archive(archive_path, files_to_archive, archive_format='zip'):
    """Создает архив из переданных файлов."""
    if archive_format == 'zip':
        with zipfile.ZipFile(archive_path, 'w') as zipf:
            for file in files_to_archive:
                zipf.write(file, arcname=os.path.basename(file))
    elif archive_format == 'rar':
        # Для rar нужно, чтобы система могла вызывать unrar/rar
        with rarfile.RarFile(archive_path, 'w') as rarf:
            for file in files_to_archive:
                rarf.write(file, arcname=os.path.basename(file))
    elif archive_format == '7z':
        with py7zr.SevenZipFile(archive_path, 'w') as szf:
            for file in files_to_archive:
                szf.write(file, arcname=os.path.basename(file))

# Генерируем файлы
for i in range(3):
    generate_text_file(os.path.join(storage_path, f"doc_{i}.txt"))
    generate_docx_file(os.path.join(storage_path, f"doc_{i}.docx"))
    generate_xlsx_file(os.path.join(storage_path, f"data_{i}.xlsx"))
    generate_pdf_file(os.path.join(storage_path, f"doc_{i}.pdf"))

# Генерируем архивы
txt_files = [os.path.join(storage_path, f"doc_{i}.txt") for i in range(2)]
docx_files = [os.path.join(storage_path, f"doc_{i}.docx") for i in range(2)]

create_archive(os.path.join(storage_path, 'archive.zip'), txt_files, 'zip')
create_archive(os.path.join(storage_path, 'archive.rar'), docx_files, 'rar')
create_archive(os.path.join(storage_path, 'archive.7z'), txt_files + docx_files, '7z')

print("Тестовое хранилище сгенерировано.")

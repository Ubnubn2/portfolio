import os
import random
from faker import Faker
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF
import zipfile
import rarfile
import py7zr

# Используем английскую локаль — текст будет на латинице
fake = Faker('en_US')

# Создаём корневую папку для тестовых данных
storage_path = 'test_storage'
if not os.path.exists(storage_path):
    os.makedirs(storage_path)

def generate_text_file(path, content=None):
    """Генерирует простой текстовый файл."""
    if content is None:
        content = fake.text()
    with open(path, 'w', encoding='utf-8') as f:
        f.write(content)

def generate_docx_file(path):
    """Генерирует .docx файл с абзацами."""
    doc = Document()
    doc.add_heading(fake.sentence(), level=1)
    for _ in range(random.randint(2, 5)):
        doc.add_paragraph(fake.paragraph())
    doc.save(path)

def generate_xlsx_file(path):
    """Генерирует .xlsx файл с таблицей."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    # Заголовки на английском
    ws.append(["First Name", "Last Name", "Email", "Phone"])
    for _ in range(10):
        ws.append([fake.first_name(), fake.last_name(), fake.email(), fake.phone_number()])
    wb.save(path)

def generate_pdf_file(path):
    """Генерирует простой .pdf файл с латинским текстом."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)          # Helvetica поддерживает латиницу
    pdf.multi_cell(200, 10, text=fake.text())   # параметр text вместо устаревшего txt
    pdf.output(path)

def create_archive(archive_path, files_to_archive, archive_format='zip'):
    """Создаёт архив из переданных файлов."""
    if archive_format == 'zip':
        with zipfile.ZipFile(archive_path, 'w') as zipf:
            for file in files_to_archive:
                zipf.write(file, arcname=os.path.basename(file))
    elif archive_format == 'rar':
        with rarfile.RarFile(archive_path, 'w') as rarf:
            for file in files_to_archive:
                rarf.write(file, arcname=os.path.basename(file))
    elif archive_format == '7z':
        with py7zr.SevenZipFile(archive_path, 'w') as szf:
            for file in files_to_archive:
                szf.write(file, arcname=os.path.basename(file))

# Генерируем отдельные файлы
for i in range(3):
    generate_text_file(os.path.join(storage_path, f"doc_{i}.txt"))
    generate_docx_file(os.path.join(storage_path, f"doc_{i}.docx"))
    generate_xlsx_file(os.path.join(storage_path, f"data_{i}.xlsx"))
    generate_pdf_file(os.path.join(storage_path, f"doc_{i}.pdf"))

# Готовим списки файлов для архивов
txt_files = [os.path.join(storage_path, f"doc_{i}.txt") for i in range(2)]
docx_files = [os.path.join(storage_path, f"doc_{i}.docx") for i in range(2)]

# Создаём архивы разных форматов
create_archive(os.path.join(storage_path, 'archive.zip'), txt_files, 'zip')
create_archive(os.path.join(storage_path, 'archive.rar'), docx_files, 'rar')
create_archive(os.path.join(storage_path, 'archive.7z'), txt_files + docx_files, '7z')

print("✅ Тестовое хранилище успешно сгенерировано!")